GOOGLE_API_KEY = "YOUR GOOGLE API"
from flask import Flask, render_template, request, jsonify, session
from flask_cors import CORS  # Add this import for CORS support
import os
import uuid
from datetime import datetime
import io
import tempfile
import logging
import json
from google import genai
from nltk.tokenize import sent_tokenize
import re
import numpy as np
from PIL import Image
from docx import Document
import PyPDF2

# Add these imports at the top with other imports
from utils.synonym_service import get_synonyms

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

# Try to import optional dependencies
try:
    import easyocr
    import cv2
    OCR_AVAILABLE = True
except ImportError:
    logging.warning("EasyOCR and/or OpenCV not available. Image OCR will be disabled.")
    OCR_AVAILABLE = False

app = Flask(__name__)
# Replace the simple CORS() call with more specific configuration
CORS(app, resources={r"/*": {"origins": "*", "allow_headers": ["Content-Type", "Authorization"], 
                             "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"]}})
app.secret_key = os.environ.get('SECRET_KEY', 'opengrammar-default-secret-key')  # For session management
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize Google AI API (Gemini)
 
 
 
client = genai.Client(api_key=GOOGLE_API_KEY)

def parse_gemini_response(response_text):
    """Parse JSON response from Gemini model, handling various formats and edge cases."""
    
    # Helper function to safely parse JSON with backslash correction
    def safe_json_loads(text):
        try:
            return json.loads(text)
        except json.JSONDecodeError as e:
            if "Invalid control character" in str(e):
                fixed_text = text.replace('\\', '\\\\')
                return json.loads(fixed_text)
            raise e

    logging.debug(f"Raw response: {response_text[:200]}...")
    
    try:
        # Try parsing using our safe loader
        return safe_json_loads(response_text)
    except json.JSONDecodeError:
        pass

    # Try to extract JSON from code blocks
    if '```' in response_text:
        parts = response_text.split('```')
        for part in parts:
            clean_part = part.strip()
            if clean_part.lower().startswith(('json', '{')):
                if (clean_part.lower().startswith('json')):
                    clean_part = clean_part[4:].strip()
                try:
                    return safe_json_loads(clean_part)
                except json.JSONDecodeError:
                    continue

    # Try to extract JSON by finding braces
    start_index = response_text.find('{')
    end_index = response_text.rfind('}')
    if (start_index != -1):
        if (end_index == -1 or not response_text.strip().endswith('}')):
            tentative = response_text[start_index:].trip() + "}"
            try:
                return safe_json_loads(tentative)
            except json.JSONDecodeError:
                pass
        elif (end_index > start_index):
            json_substr = response_text[start_index:end_index + 1]
            try:
                return safe_json_loads(json_substr)
            except json.JSONDecodeError:
                pass

    # Return error response if parsing fails
    return {
        "error": "Failed to parse response from the grammar analysis service.",
        "meta_analysis": {
            "overall_quality_score": 0,
            "confidence_level": 0,
            "summary_assessment": {
                "critical_issues": [{
                    "area": "system_error", 
                    "description": "Unable to analyze text at this time. Please try again."
                }]
            }
        }
    }
  
def check_grammar(text, goal="none", tone="none"):
    """Analyze text for grammar, style, and content improvements using the Gemini API."""
    
        
    # Sanitize input
    text = text.replace("\\", "")
    
    # Adapt goal and tone for the prompt
    goal_text = "" if goal == "none" else f"\nWriting Goal: {goal}"
    tone_text = "" if tone == "none" else f"\nTone: {tone}"
    
    # Add context-specific sections if goal or tone are specified
    context_section = ""
    if not (goal is None or goal.lower() == 'none' or tone is None or tone.lower() == 'none'):
        context_section = """
    ### 5. Purpose & Context Alignment
    - **Goal-Specific Requirements**:
    * **Academic**: Citation accuracy, thesis support, literature engagement, methodological precision
    * **Professional**: Actionability of recommendations, stakeholder consideration, compliance with industry standards
    * **Technical**: Procedural clarity, definitional precision, consistent terminology, appropriate abstraction level
    * **Persuasive**: Argument structure, emotional appeal balance, call-to-action effectiveness
    * **Instructional**: Sequencing logic, prerequisite clarity, comprehension scaffolding
    - **Tone Consistency Factors**:
    * Formality spectrum positioning
    * Stance markers and evaluation language
    * Cultural sensitivity indicators
    * Power distance acknowledgment
    * Emotional resonance calibration
    * Brand voice alignment metrics
    """
    
    # Build the core prompt
    prompt = (
        f"Act as an advanced writing analysis system combining linguistic expertise with professional editing standards. {goal_text}{tone_text}\n\n"
        "## Analysis Framework\n"
        "Perform multi-dimensional text analysis at sentence, paragraph, and document levels:\n\n"
        "### 1. Linguistic Mechanics (Micro-level)\n"
        "- **Punctuation**: Proper usage of commas, semicolons, colons, quotation marks, apostrophes, hyphens, dashes, and brackets\n"
        "- **Grammar Structure**: Subject-verb agreement, pronoun-antecedent agreement, noun-modifier placement\n"
        "- **Verb Usage**: Tense consistency, aspect appropriateness, mood selection, voice choice\n"
        "- **Articles & Determiners**: Correct a/an/the usage, quantifier appropriateness\n"
        "- **Spelling & Word Formation**: Common misspellings, compound words, contractions, possessives\n"
        "- **Syntax Patterns**: Phrase ordering, clause construction, modifier placement\n\n"
        "### 2. Stylistic Elements (Meso-level)\n"
        "- **Sentence Construction**:\n"
        "  * Variety (simple, compound, complex, compound-complex)\n"
        "  * Length distribution (short for emphasis, medium for readability, long for elaboration)\n"
        "  * Topic position and information flow\n"
        "  * Parallelism in coordinated elements\n"
        "- **Rhetorical Devices**:\n"
        "  * Emphasis techniques (fronting, clefting, etc.)\n"
        "  * Transitional expressions for coherence\n"
        "  * Repetition patterns (anaphora, epistrophe, etc.)\n"
        "- **Lexical Selection**:\n"
        "  * Register-appropriate vocabulary\n"
        "  * Connotation awareness\n"
        "  * Domain-specific terminology correctness\n"
        "  * Collocation naturalness\n"
        "  * Redundancy elimination\n"
        "- **Voice & Perspective**:\n"
        "  * Active/passive voice strategic usage\n"
        "  * Consistency in narrative perspective (1st/2nd/3rd person)\n"
        "  * Stance markers and hedging appropriateness\n\n"
        "### 3. Content Effectiveness (Macro-level)\n"
        "- **Structural Coherence**:\n"
        "  * Logical progression between sentences\n"
        "  * Paragraph unity and development\n"
        "  * Section-to-section transitions\n"
        "  * Introduction-body-conclusion relationship\n"
        "- **Argumentative Integrity**:\n"
        "  * Claim-evidence relationships\n"
        "  * Logical fallacy identification\n"
        "  * Counterargument anticipation\n"
        "  * Premise-conclusion connections\n"
        "- **Information Density**:\n"
        "  * Content-to-word ratio optimization\n"
        "  * Unnecessary qualification removal\n"
        "  * Appropriate level of detail/explanation\n"
        "- **Audience Adaptation**:\n"
        "  * Presupposition management\n"
        "  * Background knowledge assumptions\n"
        "  * Specialized terminology explanation\n"
        "  * Cultural reference accessibility\n"
        "- **Clarity Metrics**:\n"
        "  * Readability scores (Flesch-Kincaid, SMOG, Coleman-Liau)\n"
        "  * Fog index and complexity measures\n"
        "  * Sentence-to-sentence cohesion strength\n\n"
        "### 4. Common Error Patterns\n"
        "- **Syntactic Issues**:\n"
        "  * Run-on sentences (comma splices, fused sentences)\n"
        "  * Sentence fragments and incomplete structures\n"
        "  * Faulty parallelism in series or comparisons\n"
        "  * Dangling and misplaced modifiers\n"
        "  * Squinting modifiers creating ambiguity\n"
        "- **Semantic Problems**:\n"
        "  * Malapropisms and incorrect word usage\n"
        "  * Tautologies and circular reasoning\n"
        "  * Vague pronoun references\n"
        "  * Mixed metaphors and inconsistent imagery\n"
        "- **Stylistic Weaknesses**:\n"
        "  * Excessive nominalization\n"
        "  * Overuse of expletive constructions (it is, there are)\n"
        "  * Unnecessarily complex phrases\n"
        "  * Clichés and stock expressions\n"
        "  * Hedging overuse (somewhat, kind of, etc.)\n"
        "- **Formatting Inconsistencies**:\n"
        "  * Heading hierarchy violations\n"
        "  * List structure inconsistencies\n"
        "  * Quotation style variations\n"
        "  * Number/date format inconsistencies\n"
        "  * Citation style deviations\n\n"
    )
    
    # Add the context section if applicable
    prompt += context_section
    
    # Add response schema and processing instructions
    prompt += (
        "\n## Advanced Analytics\n"
        "- Compare text against corpus-based norms for genre/domain\n"
        "- Track lexical diversity (type-token ratio, unique word percentage)\n"
        "- Identify overused words and phrases with frequency analysis\n"
        "- Measure sentence-initial variety (SIV) score\n"
        "- Calculate coherence metrics through lexical chain analysis\n"
        "- Evaluate passive construction percentage against industry benchmarks\n\n"
        "## Response Schema\n"
        "Return a precisely structured JSON object:\n\n"
        "```json\n"
        "{\n"
        '  "meta_analysis": {\n'
        '    "overall_quality_score": <integer 0-100>,\n'
        '    "confidence_level": <float 0.0-1.0 indicating analysis reliability>,\n'
        '    "document_metrics": {\n'
        '      "character_count": <total character count including spaces>,\n'
        '      "word_count": <total word count>,\n'
        '      "sentence_count": <total sentence count>,\n'
        '      "paragraph_count": <total paragraph count>,\n'
        '      "average_words_per_sentence": <float>,\n'
        '      "average_sentences_per_paragraph": <float>,\n'
        '      "readability_scores": {\n'
        '        "flesch_kincaid_grade": <float>,\n'
        '        "flesch_reading_ease": <float>,\n'
        '        "gunning_fog": <float>\n'
        '      },\n'
        '      "lexical_diversity": <type-token ratio as float>,\n'
        '      "passive_voice_percentage": <float>\n'
        '    },\n'
        '    "summary_assessment": {\n'
        '      "major_strengths": [\n'
        '        {"area": "<strength category>", "description": "<specific strength>"}\n'
        '      ],\n'
        '      "critical_issues": [\n'
        '        {"area": "<issue category>", "description": "<specific issue>", "frequency": <count>}\n'
        '      ],\n'
        '      "improvement_priorities": [\n'
        '        {"recommendation": "<actionable suggestion>", "impact_level": "<high|medium|low>"}\n'
        '      ]\n'
        '    }\n'
        '  },\n'
        '  "sentence_analysis": [\n'
        '    {\n'
        '      "id": <sequential integer>,\n'
        '      "original_text": "<exact original sentence>",\n'
        '      "improved_text": "<enhanced version or \'NO_REVISION_NEEDED\'>",\n'
        '      "position": {\n'
        '        "start_char": <integer position>,\n'
        '        "end_char": <integer position>,\n'
        '        "paragraph_number": <integer>\n'
        '      },\n'
        '      "metrics": {\n'
        '        "length_characters": <integer>,\n'
        '        "length_words": <integer>,\n'
        '        "complexity_score": <float 0.0-1.0>,\n'
        '        "revision_impact": <float representing improvement percentage>\n'
        '      },\n'
        '      "identified_issues": [\n'
        '        {\n'
        '          "category": "<grammar|style|clarity|logic|tone|other>",\n'
        '          "subcategory": "<specific issue type>",\n'
        '          "severity": "<critical|major|minor>",\n'
        '          "explanation": "<detailed issue description>",\n'
        '          "location": {\n'
        '            "start_char": <integer relative to sentence>,\n'
        '            "end_char": <integer relative to sentence>\n'
        '          },\n'
        '          "correction_rationale": "<explanation of why the change improves the text>"\n'
        '        }\n'
        '      ],\n'
        '      "improvement_status": "<perfect|revised|needs_attention>",\n'
        '      "context_notes": "<observations about sentence role in paragraph/document>"\n'
        '    }\n'
        '  ],\n'
        '  "paragraph_analysis": [\n'
        '    {\n'
        '      "id": <sequential integer>,\n'
        '      "position": {"start_char": <integer>, "end_char": <integer>},\n'
        '      "structure_assessment": {\n'
        '        "topic_sentence_strength": <float 0.0-1.0>,\n'
        '        "development_quality": <float 0.0-1.0>,\n'
        '        "unity_score": <float 0.0-1.0>,\n'
        '        "transition_effectiveness": <float 0.0-1.0>\n'
        '      },\n'
        '      "improvement_suggestions": "<paragraph-level recommendations>"\n'
        '    }\n'
        '  ],\n'
        '  "document_coherence": {\n'
        '    "global_structure": "<assessment of overall organization>",\n'
        '    "thematic_consistency": <float 0.0-1.0>,\n'
        '    "logical_flow_rating": <float 0.0-1.0>,\n'
        '    "structural_recommendations": "<document-level organization suggestions>"\n'
        '  }\n'
        '}\n'
        '```\n\n'
        "## Processing Instructions\n"
        "1. Prioritize impactful corrections that enhance clarity, precision, and effectiveness\n"
        "2. Preserve the writer's unique voice and stylistic choices when they serve the communication purpose\n"
        "3. Consider context before applying prescriptive rules (some \"violations\" may be intentional for effect)\n"
        "4. Provide educational explanations that help improve writing skills, not just fix current text\n"
        "5. Balance concision with comprehensiveness in feedback\n"
        "6. Adapt analysis depth to text length (more detailed for shorter texts, more strategic for longer ones)\n"
        "7. When goal/tone are specified, weight recommendations accordingly\n"
        "8. Ensure all scores are calibrated to professional editing standards in the relevant domain\n\n"
        f"Text to analyze:\n{text}"
    )
    
    # Split text into sentences with position tracking
    current_pos = 0
    sentences_data = []
    
    # Extract sentences and track their positions
    for paragraph_num, paragraph in enumerate(text.split('\n\n')):
        if not paragraph.strip():
            continue
            
        paragraph_start = text.find(paragraph, current_pos)
        if paragraph_start == -1:  # Safety check
            continue
            
        paragraph_sentences = sent_tokenize(paragraph)
        paragraph_pos = paragraph_start
        
        for sentence in paragraph_sentences:
            sentence_start = text.find(sentence, paragraph_pos)
            if sentence_start == -1:  # Safety check
                continue
                
            sentence_end = sentence_start + len(sentence)
            sentences_data.append({
                "text": sentence,
                "start": sentence_start,
                "end": sentence_end,
                "paragraph": paragraph_num + 1
            })
            paragraph_pos = sentence_end
            
        current_pos = paragraph_start + len(paragraph)

    # Add sentence position data to the prompt
    prompt += "\n\nText with positions:\n"
    for i, sent in enumerate(sentences_data):
        prompt += f"[{sent['start']}-{sent['end']}]: {sent['text']}\n"
    
    try:
        # Call Gemini API to analyze the text
       
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt,
             
        )
        
        analysis = parse_gemini_response(response.text)
        return analysis
        
    except Exception as e:
        logging.error(f"Error during grammar check: {str(e)}")
        return {
            "error": f"Failed to analyze text: {str(e)}",
            "meta_analysis": {
                "overall_quality_score": 0,
                "confidence_level": 0,
                "summary_assessment": {
                    "critical_issues": [{
                        "area": "system_error", 
                        "description": "An error occurred during analysis."
                    }]
                }
            }
        }

def extract_text_from_image(file_bytes):
    """Extract text from image files using EasyOCR"""
    if not OCR_AVAILABLE:
        return "OCR functionality is not available. Please install easyocr and opencv-python packages."
        
    try:
        npimg = np.frombuffer(file_bytes, np.uint8)
        img = cv2.imdecode(npimg, cv2.IMREAD_COLOR)
        reader = easyocr.Reader(['en'])
        result = reader.readtext(img, detail=0, paragraph=True)
        return " ".join(result)
    except Exception as e:
        logging.error(f"Error during OCR processing: {str(e)}")
        return f"Error extracting text from image: {str(e)}"

def extract_text_from_pdf(file_bytes):
    """Extract text from PDF files using PyPDF2"""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
            else:
                logging.warning(f"No text extracted from page {page_num+1}")
                
        if not text.strip():
            return "No text could be extracted from the PDF. It may be scanned or contain only images."
            
        return text
    except Exception as e:
        logging.error(f"Error during PDF processing: {str(e)}")
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(file_bytes):
    """Extract text from DOCX files using python-docx"""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += "\n"
                
        if not text.strip():
            return "No text could be extracted from the document."
            
        return text
    except Exception as e:
        logging.error(f"Error during DOCX processing: {str(e)}")
        return f"Error extracting text from DOCX: {str(e)}"

@app.route('/')
def index():
    """Render the main grammar checker page"""
    return render_template('grammar.html')  

@app.route('/check_grammar', methods=['POST'])
def check_grammar_route():
    """API endpoint to check grammar and analyze text"""
    # Validate request data
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400
        
    text = data.get('text', '').strip()
    if not text:
        return jsonify({'error': 'Empty text provided'}), 400
        
    # Extract optional parameters with defaults
    goal = data.get('goal', 'none')
    tone = data.get('tone', 'none')
    
    # Set a reasonable limit for text length
    MAX_TEXT_LENGTH = 10000  # Adjust as needed based on API limits
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH]
        truncated = True
    else:
        truncated = False
    
    # Call the grammar checking function
    result = check_grammar(text, goal, tone)
    
    # Add truncation warning if applicable
    if truncated and 'meta_analysis' in result:
        if 'summary_assessment' not in result['meta_analysis']:
            result['meta_analysis']['summary_assessment'] = {}
            
        if 'critical_issues' not in result['meta_analysis']['summary_assessment']:
            result['meta_analysis']['summary_assessment']['critical_issues'] = []
            
        result['meta_analysis']['summary_assessment']['critical_issues'].append({
            "area": "input_length",
            "description": f"Text was truncated to {MAX_TEXT_LENGTH} characters due to length limitations."
        })
    
    logging.info(f"Grammar check completed: score={result.get('meta_analysis', {}).get('overall_quality_score', 'N/A')}")
    return jsonify(result)

@app.route('/upload_file', methods=['POST'])
def upload_file():
    """API endpoint to handle file uploads and extract text"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    # Check file size
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    file_bytes = file.read()
    if len(file_bytes) > MAX_FILE_SIZE:
        return jsonify({'error': 'File size exceeds the 10MB limit'}), 413

    # Process based on file extension
    try:
        file_extension = file.filename.lower().split('.')[-1]
        
        if file_extension in ['pdf']:
            text = extract_text_from_pdf(file_bytes)
        elif file_extension in ['docx', 'doc']:
            text = extract_text_from_docx(file_bytes)
        elif file_extension in ['png', 'jpg', 'jpeg']:
            if not OCR_AVAILABLE:
                return jsonify({'error': 'Image processing is not available. Required packages not installed.'}), 501
            text = extract_text_from_image(file_bytes)
        else:
            return jsonify({'error': f'Unsupported file type: .{file_extension}'}), 415

        # Validate the extracted text
        if not text or not text.strip():
            return jsonify({'error': 'No text could be extracted from the file'}), 422
            
        if text.startswith('Error extracting text'):
            return jsonify({'error': text}), 422

        # Limit text length if necessary
        MAX_TEXT_LENGTH = 10000  # Same as in check_grammar_route
        if len(text) > MAX_TEXT_LENGTH:
            text = text[:MAX_TEXT_LENGTH]
            return jsonify({
                'text': text,
                'warning': f'Text was truncated to {MAX_TEXT_LENGTH} characters due to length limitations.'
            })

        return jsonify({'text': text})
        
    except Exception as e:
        logging.error(f"File upload error: {str(e)}")
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

@app.route('/get_synonyms/<word>')
def get_word_synonyms(word):
    """API endpoint to get synonyms for a word"""
    try:
        synonyms = get_synonyms(word)
        return jsonify({'synonyms': synonyms})
    except Exception as e:
        logging.error(f"Error getting synonyms for {word}: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.errorhandler(413)
def request_entity_too_large(error):
    """Custom error handler for large file uploads"""
    return jsonify({'error': 'File size exceeds the maximum limit'}), 413

@app.errorhandler(500)
def internal_server_error(error):
    """Custom error handler for server errors"""
    return jsonify({'error': 'Internal server error occurred. Please try again later.'}), 500

# Add route for static files if needed
@app.route('/static/<path:filename>')
def serve_static(filename):
    return app.send_static_file(filename)

# Add explicit OPTIONS handler for preflight requests
@app.route('/', methods=['OPTIONS'])
@app.route('/<path:path>', methods=['OPTIONS'])
def options_handler(path=None):
    return '', 204

if __name__ == '__main__':
    # Ensure required directories exist
    os.makedirs('static', exist_ok=True)
    
    # Start the Flask application
    port = int(os.environ.get('PORT', 5001))
    app.run(host='0.0.0.0', port=port, debug=os.environ.get('FLASK_ENV') == 'development')
